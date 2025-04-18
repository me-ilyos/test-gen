import os
import tempfile
import pytest
from docx import Document
from formatters import (
    transform_to_student_format,
    transform_to_program_format,
    create_word_document
)


@pytest.fixture
def sample_questions():
    """Create sample questions data structure."""
    return {
        "questions": [
            {
                "id": 1,
                "text": "What is Python?",
                "variants": [
                    {"id": 1, "text": "A snake"},
                    {"id": 2, "text": "A programming language"},
                    {"id": 3, "text": "An operating system"},
                    {"id": 4, "text": "A database"}
                ],
                "correct": 2
            },
            {
                "id": 2,
                "text": "What does CPU stand for?",
                "variants": [
                    {"id": 1, "text": "Central Processing Unit"},
                    {"id": 2, "text": "Computer Processing Unit"},
                    {"id": 3, "text": "Central Program Utility"},
                    {"id": 4, "text": "Computer Program Utility"}
                ],
                "correct": 1
            }
        ]
    }


def test_transform_to_student_format_with_variants(sample_questions):
    """Test converting questions to student format with answer variants."""
    result = transform_to_student_format(sample_questions, include_variants=True)
    lines = result.strip().split("\n")
    
    # Check first question
    assert lines[0] == "1. What is Python?"
    assert lines[1] == "a) A snake"
    assert lines[2] == "b) A programming language"
    assert lines[3] == "c) An operating system"
    assert lines[4] == "d) A database"
    
    # Check second question (accounting for the blank line between questions)
    assert lines[6] == "2. What does CPU stand for?"
    assert lines[7] == "a) Central Processing Unit"
    assert lines[8] == "b) Computer Processing Unit"
    assert lines[9] == "c) Central Program Utility"
    assert lines[10] == "d) Computer Program Utility"


def test_transform_to_student_format_without_variants(sample_questions):
    """Test converting questions to student format without answer variants."""
    result = transform_to_student_format(sample_questions, include_variants=False)
    lines = result.strip().split("\n")
    
    # Check questions only (no variants)
    assert lines[0] == "1. What is Python?"
    assert lines[2] == "2. What does CPU stand for?"
    
    # Make sure no variants are included
    assert "a) A snake" not in result
    assert "a) Central Processing Unit" not in result


def test_transform_to_program_format(sample_questions):
    """Test converting questions to program format (HEMIS)."""
    result = transform_to_program_format(sample_questions)
    lines = result.strip().split("\n")
    
    # Check first question format
    assert lines[0] == "What is Python?"
    assert lines[1] == "===="
    assert lines[2] == "A snake"
    assert lines[3] == "===="
    assert lines[4] == "#A programming language"  # Correct answer with # marker
    assert lines[5] == "===="
    
    # Check separation between questions
    assert "+++++" in result or "++++" in result
    
    # Check second question has correct answer marked
    cpu_index = lines.index("What does CPU stand for?")
    assert "#Central Processing Unit" in lines[cpu_index:cpu_index+10]


def test_create_word_document(sample_questions):
    """Test creating a Word document with questions and answers."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp:
        temp_path = temp.name
    
    # Create the document
    create_word_document(sample_questions, temp_path)
    
    # Open and verify the document
    doc = Document(temp_path)
    
    # Count tables (should be one per question)
    assert len(doc.tables) == 2
    
    # Check first question
    table1 = doc.tables[0]
    assert "What is Python?" in table1.cell(0, 0).text
    assert "A programming language" in table1.cell(1, 0).text  # Correct answer first
    
    # Check second question
    table2 = doc.tables[1]
    assert "What does CPU stand for?" in table2.cell(0, 0).text
    assert "Central Processing Unit" in table2.cell(1, 0).text  # Correct answer first
    
    # Clean up
    os.unlink(temp_path)