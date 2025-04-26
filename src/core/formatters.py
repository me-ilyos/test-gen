"""
Formatters for converting questions to different output formats.

This module contains functions to transform question data into various
output formats including student format, HEMIS format, and Word documents.
"""

from docx import Document
from typing import Dict, List, Optional


def transform_to_student_format(json_data: Dict, include_variants: bool = True) -> str:
    """
    Convert questions to student format with or without answer variants.

    Args:
        json_data: Dictionary containing questions data
        include_variants: Whether to include answer variants in output

    Returns:
        Formatted text for student use
    """
    output = []
    questions = json_data["questions"]

    for question in questions:
        # Add question
        output.append(f"{question['id']}. {question['text']}")

        # Add variants if requested
        if include_variants:
            for variant in question["variants"]:
                letter = chr(96 + variant["id"])  # Convert 1->a, 2->b, etc.
                output.append(f"{letter}) {variant['text']}")

        # Add blank line between questions
        output.append("")

    return "\n".join(output)


def transform_to_program_format(json_data: Dict) -> str:
    """
    Convert questions to HEMIS format with markers for correct answers.

    Args:
        json_data: Dictionary containing questions data

    Returns:
        Formatted text in HEMIS format
    """
    output = []
    questions = json_data["questions"]

    for i, question in enumerate(questions):
        # Add question
        output.append(question["text"])
        output.append("====")

        # Add variants
        for variant in question["variants"]:
            is_correct = variant["id"] == question["correct"]
            marker = "#" if is_correct else ""
            output.append(f"{marker}{variant['text']}")
            output.append("====")

        # Add separator between questions (except after the last one)
        if i < len(questions) - 1:
            output.append("++++")

    return "\n".join(output)

def create_word_document(json_data: Dict, output_path: str) -> None:
    """
    Create a Word document with questions in tables.
    First row contains the question, second row contains the correct answer,
    and remaining rows contain incorrect answers.
    """
    doc = Document()
    questions = json_data["questions"]

    for question in questions:
        # Create table for this question
        table = doc.add_table(rows=len(question["variants"]) + 1, cols=1)
        table.style = "Table Grid"

        # Add question to first row
        table.cell(0, 0).text = question["text"]

        # Find the correct answer
        correct_variant = None
        incorrect_variants = []
        
        for variant in question["variants"]:
            if variant["id"] == question["correct"]:
                correct_variant = variant
            else:
                incorrect_variants.append(variant)
        
        # Add correct answer to second row
        if correct_variant:
            table.cell(1, 0).text = correct_variant["text"]
            
            # Add incorrect answers to remaining rows
            for i, variant in enumerate(incorrect_variants):
                table.cell(i + 2, 0).text = variant["text"]
        else:
            # If no correct answer is found, just add all variants in order
            for i, variant in enumerate(question["variants"]):
                table.cell(i + 1, 0).text = variant["text"]

        # Add space between questions
        doc.add_paragraph()

    # Save the document
    doc.save(output_path)



def create_student_word_document(json_data: Dict, output_path: str, include_variants: bool = True) -> None:
    """Create a Word document with questions in student format."""
    doc = Document()
    questions = json_data["questions"]

    for question in questions:
        doc.add_paragraph(f"{question['id']}. {question['text']}")

        if include_variants:
            for variant in question["variants"]:
                letter = chr(96 + variant["id"])
                doc.add_paragraph(f"{letter}) {variant['text']}")

        doc.add_paragraph()  # Space between questions

    doc.save(output_path)
