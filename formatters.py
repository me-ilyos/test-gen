from docx import Document


def transform_to_student_format(json_data, include_variants=True):
    output = []
    for question in json_data["questions"]:
        output.append(f"{question['id']}. {question['text']}")

        if include_variants:
            for variant in question["variants"]:
                prefix = chr(96 + variant["id"])
                output.append(f"{prefix}) {variant['text']}")

        output.append("")

    return "\n".join(output)


def transform_to_program_format(json_data):
    output = []
    questions = json_data["questions"]

    for i, question in enumerate(questions):
        output.append(f"{question['text']}")
        output.append("====")

        for variant in question["variants"]:
            is_correct = variant["id"] == question["correct"]
            marker = "#" if is_correct else ""
            output.append(f"{marker}{variant['text']}")
            output.append("====")

        if i < len(questions) - 1:
            output.append("++++")

    return "\n".join(output)


def create_word_document(json_data, output_path):
    doc = Document()

    for question in json_data["questions"]:
        num_variants = len(question["variants"])
        table = doc.add_table(rows=num_variants + 1, cols=1)
        table.style = "Table Grid"

        question_cell = table.cell(0, 0)
        question_cell.text = question["text"]

        correct_variant = next(
            v for v in question["variants"] if v["id"] == question["correct"]
        )

        table.cell(1, 0).text = correct_variant["text"]

        row = 2
        for variant in question["variants"]:
            if variant["id"] != question["correct"]:
                table.cell(row, 0).text = variant["text"]
                row += 1

        doc.add_paragraph()

    doc.save(output_path)
